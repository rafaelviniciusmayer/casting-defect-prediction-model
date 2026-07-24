"""
Atualiza os arquivos Word do artigo com os resultados das Fases 1-3 (ajustes da banca).

Arquivo final único (artigo principal):
  Artigo 2 - FINAL - BANCA REVISADO.docx

Material suplementar (documento separado exigido pela submissão):
  Artigo 2 - Suplementar - FINAL - BANCA REVISADO.docx

Os originais (Artigo 2 - FINAL.docx) nunca são alterados.
Versões intermediárias antigas (- TABELAS, _table_fix_test, etc.) são removidas ao final.
"""

from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.text.paragraph import Paragraph

REPO = Path(__file__).resolve().parent
FIGURES = REPO / "figures"

BASE = Path(r"c:\Users\Cliente\Desktop\RETA FINAL\REVISADO\FEITO\SUBMETER\2")
MAIN_SRC = BASE / "Artigo 2 - FINAL.docx"
SUPP_SRC = BASE / "Artigo 2 - Suplementar - FINAL.docx"
MAIN_OUT = BASE / "Artigo 2 - FINAL - BANCA REVISADO.docx"
SUPP_OUT = BASE / "Artigo 2 - Suplementar - FINAL - BANCA REVISADO.docx"

# Métricas oficiais (pipeline executado — Fase 3, F-beta β=2, NN, teste 20%)
OFFICIAL = {
    "recall_micro": "0.9828",
    "precision_micro": "0.5068",
    "f1_micro": "0.6687",
    "f1_macro": "0.6150",
}
# Table 8: comparação entre modelos com thresholds recall-first (Fase 2)
TABLE8_NN = {
    "recall": "0.9930",
    "f1_micro": "0.6639",
    "inference_ms": "0.30",
}


def find_para(doc: Document, startswith: str) -> Paragraph:
    for p in doc.paragraphs:
        if p.text.strip().startswith(startswith):
            return p
    raise ValueError(f"Parágrafo não encontrado: {startswith!r}")


def find_para_contains(doc: Document, snippet: str) -> Paragraph:
    for p in doc.paragraphs:
        if snippet in p.text:
            return p
    raise ValueError(f"Parágrafo não encontrado (contém): {snippet!r}")


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style:
        new_para.style = style
    return new_para


def insert_paragraph_before(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style:
        new_para.style = style
    return new_para


MATH_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
MATH_TAGS = {f"{{{MATH_NS}}}oMath", f"{{{MATH_NS}}}oMathPara"}


def paragraph_has_math(paragraph: Paragraph) -> bool:
    """True if paragraph contains Word equation objects (must not be overwritten)."""
    for el in paragraph._p.iter():
        if el.tag in MATH_TAGS:
            return True
    return False


def cell_has_math(cell) -> bool:
    for el in cell._tc.iter():
        if el.tag in MATH_TAGS:
            return True
    return False


def set_paragraph_text(paragraph: Paragraph, text: str) -> bool:
    """Replace plain-text paragraph content. Skips paragraphs with equations."""
    if paragraph_has_math(paragraph):
        return False
    p_el = paragraph._p
    for child in list(p_el):
        tag = child.tag.split("}")[-1]
        if tag in ("r", "hyperlink", "ins", "del"):
            p_el.remove(child)
    paragraph.add_run(text)
    return True


def insert_note_after(paragraph: Paragraph, text: str) -> Paragraph:
    """Add a revision note after a paragraph without altering it (preserves equations)."""
    return insert_paragraph_after(paragraph, text, "New paragraph")


def replace_if_contains(doc: Document, snippet: str, new_text: str) -> bool:
    for p in doc.paragraphs:
        if snippet in p.text:
            if paragraph_has_math(p):
                insert_note_after(p, new_text)
                return True
            set_paragraph_text(p, new_text)
            return True
    return False


def set_cell_value(cell, value: str) -> None:
    """Update a table cell while preserving equation objects if present."""
    if cell_has_math(cell):
        return
    text = str(value)
    if cell.paragraphs:
        para = cell.paragraphs[0]
        if paragraph_has_math(para):
            return
        for extra in cell.paragraphs[1:]:
            extra._p.getparent().remove(extra._p)
        p_el = para._p
        for child in list(p_el):
            if child.tag.split("}")[-1] == "r":
                p_el.remove(child)
        para.add_run(text)
    else:
        cell.text = text


def fill_table(table, rows: list[list[str]]) -> None:
    """Fill table via grid coordinates; skips cells that contain equations."""
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    while len(table.rows) < len(rows):
        table.add_row()
    for ri, row_data in enumerate(rows):
        for ci in range(n_cols):
            val = row_data[ci] if ci < len(row_data) else ""
            cell = table.cell(ri, ci)
            if not cell_has_math(cell):
                set_cell_value(cell, val)
    for ri in range(len(rows), len(table.rows)):
        for ci in range(n_cols):
            cell = table.cell(ri, ci)
            if not cell_has_math(cell):
                set_cell_value(cell, "")


def update_table_18(table) -> None:
    fill_table(
        table,
        [
            ["Model", "F1-Micro", "F1-Macro", "Precision", "Recall", "Training (s)", "Inference (ms/100)"],
            ["PyTorch NN", "0.6639", "0.6088", "0.4986", "0.9930", "48.5", "0.30"],
            ["XGBoost", "0.6610", "0.6200", "0.4984", "0.9811", "23.4", "16.96"],
            ["Random Forest", "0.5847", "0.5600", "0.4137", "0.9968", "55.8", "1269.45"],
            ["Logistic Regression (L2)", "0.6630", "0.6206", "0.4999", "0.9844", "11.3", "2.85"],
            ["Logistic Regression (L1)", "0.6636", "0.6191", "0.5005", "0.9844", "510.2", "2.77"],
        ],
    )


def update_table_17_smote_reference(table) -> None:
    """Table 9 — frequências originais (referência do ablation; SMOTE não usado no pipeline final)."""
    fill_table(
        table,
        [
            [
                "Defect Type",
                "Original Count",
                "Original Freq. (%)",
                "After SMOTE",
                "After SMOTE Freq. (%)",
                "Ratio Before → After",
            ],
            ["Gas Porosity", "644", "3.22%", "15,650", "36.95%", "1:30 → 1:1.7"],
            ["Density Deviation", "596", "2.98%", "14,450", "34.12%", "1:33 → 1:1.9"],
            ["Cold Shut", "580", "2.90%", "14,050", "33.18%", "1:34 → 1:2.0"],
            ["Gas Bubbles", "520", "2.60%", "12,600", "29.75%", "1:38 → 1:2.3"],
            ["Incomplete Fill", "420", "2.10%", "10,200", "24.08%", "1:47 → 1:2.9"],
            ["Surface Blisters", "380", "1.90%", "9,200", "21.72%", "1:52 → 1:3.2"],
            ["Flow Lines", "213", "1.06%", "5,399", "12.75%", "1:93 → 1:6.8"],
            ["Shrinkage Porosity", "73", "0.36%", "1,653", "3.90%", "1:273 → 1:25"],
        ],
    )


def update_table_19_thresholds(table) -> None:
    """Table 11 — thresholds F-beta (β=2), modelo neural oficial."""
    fill_table(
        table,
        [
            ["Defect Type", "Optimal Threshold", "Recall", "Precision", "F1-Score"],
            ["Gas Porosity", "0.05", "1.000", "0.540", "0.701"],
            ["Density Deviation", "0.10", "1.000", "0.596", "0.747"],
            ["Cold Shut", "0.21", "1.000", "0.590", "0.742"],
            ["Gas Bubbles", "0.42", "0.977", "0.624", "0.761"],
            ["Incomplete Fill", "0.31", "0.981", "0.520", "0.680"],
            ["Shrinkage Porosity", "0.54", "1.000", "0.500", "0.667"],
            ["Low Tensile Strength", "0.56", "0.963", "0.542", "0.694"],
            ["Warpage", "0.53", "1.000", "0.343", "0.510"],
            ["Ejector Pin Marks", "0.50", "0.000", "0.000", "—"],
        ],
    )


def update_table_20_confusion(table) -> None:
    """Table 12 — matriz de confusão (8 defeitos mais frequentes + micro-avg)."""
    fill_table(
        table,
        [
            ["Defect Type", "TP", "TN", "FP", "FN", "Recall", "Precision"],
            ["Gas Porosity", "141", "4,739", "120", "0", "1.000", "0.540"],
            ["Density Deviation", "140", "4,765", "95", "0", "1.000", "0.596"],
            ["Cold Shut", "138", "4,766", "96", "0", "1.000", "0.590"],
            ["Gas Bubbles", "126", "4,795", "76", "3", "0.977", "0.624"],
            ["Incomplete Fill", "103", "4,800", "95", "2", "0.981", "0.520"],
            ["Surface Blisters", "91", "4,788", "119", "2", "0.978", "0.433"],
            ["Blisters Post Treatment", "87", "4,840", "68", "5", "0.946", "0.561"],
            ["Low Elongation", "83", "4,843", "73", "1", "0.988", "0.532"],
            ["Overall (micro-avg)", "1,824", "136,369", "1,775", "32", "0.983", "0.507"],
        ],
    )


# Table numbering in the manuscript (sequential 1–10; equation display tables are not numbered)
TABLE_SMOTE = 7
TABLE_MODEL_COMP = 8
TABLE_THRESHOLDS = 9
TABLE_CONFUSION = 10

# Figure numbering in the manuscript (sequential 1–11)
FIG_WORKFLOW = 1
FIG_FEAT_ENG = 2
FIG_THRESHOLD = 3
FIG_DASHBOARD = 4
FIG_CONFUSION = 5
FIG_LEARNING = 6
FIG_IMPORTANCE = 7
FIG_IMPORTANCE_PHASE = 8
FIG_ROC_TOP = 9
FIG_ROC_MICRO = 10
FIG_CALIBRATION = 11

THRESHOLD_CAPTION = (
    f"Figure {FIG_THRESHOLD}. Threshold strategy comparison "
    "(recall-first vs. F-beta β=2) for the neural network."
)


def replace_in_paragraph(paragraph: Paragraph, replacements: list[tuple[str, str]]) -> bool:
    """Apply text replacements if paragraph has no equations."""
    if paragraph_has_math(paragraph) or not paragraph.text:
        return False
    new_text = paragraph.text
    for old, new in replacements:
        new_text = new_text.replace(old, new)
    if new_text == paragraph.text:
        return False
    set_paragraph_text(paragraph, new_text)
    return True


def replace_in_all_paragraphs(doc: Document, replacements: list[tuple[str, str]]) -> int:
    count = 0
    for p in doc.paragraphs:
        if replace_in_paragraph(p, replacements):
            count += 1
    return count


def fix_table_numbering(doc: Document) -> None:
    """
    Renumber results tables to sequential order 1–10.

    Original manuscript skipped 7 and duplicated 9:
      SMOTE frequencies (was 9) -> 7
      Model comparison (8)       -> 8
      Thresholds (was 11/9)      -> 9
      Confusion matrix (was 12)  -> 10
    """
    # Step 1: SMOTE-specific Table 9 -> 7 (before any generic renumbering)
    smote_replacements = [
        (
            "Table 9. Defect Frequencies and SMOTE Balancing Effect",
            f"Table {TABLE_SMOTE}. Defect Frequencies and SMOTE Balancing Effect",
        ),
        (
            "Table 9 presents the original frequency",
            f"Table {TABLE_SMOTE} presents the original frequency",
        ),
        (
            "Table 9 illustrates the dramatic effect of SMOTE",
            f"Table {TABLE_SMOTE} illustrates the dramatic effect of SMOTE",
        ),
        (
            "Table 9 retains the original SMOTE frequency projections",
            f"Table {TABLE_SMOTE} retains the original SMOTE frequency projections",
        ),
    ]
    replace_in_all_paragraphs(doc, smote_replacements)

    # Step 2: Thresholds Table 11 -> 9
    replace_in_all_paragraphs(
        doc,
        [
            (
                "Table 11. Optimized Thresholds and Performance for Selected Defects",
                f"Table {TABLE_THRESHOLDS}. Optimized Thresholds and Performance for Selected Defects",
            ),
            ("Table 11 presents", f"Table {TABLE_THRESHOLDS} presents"),
        ],
    )

    # Step 3: Confusion Table 12 -> 10
    replace_in_all_paragraphs(
        doc,
        [
            (
                "Table 12. Confusion Matrix Metrics",
                f"Table {TABLE_CONFUSION}. Confusion Matrix Metrics",
            ),
            ("Table 12 presents", f"Table {TABLE_CONFUSION} presents"),
            ("row in Table 12", f"row in Table {TABLE_CONFUSION}"),
            ("in Table 12 ", f"in Table {TABLE_CONFUSION} "),
        ],
    )

    # Step 4: Residual duplicates from original (thresholds also labeled Table 9)
    replace_in_all_paragraphs(
        doc,
        [
            (
                "Table 9. Optimized Thresholds and Performance for Selected Defects",
                f"Table {TABLE_THRESHOLDS}. Optimized Thresholds and Performance for Selected Defects",
            ),
        ],
    )


def remove_paragraphs_matching(doc: Document, predicate) -> int:
    """Remove body paragraphs for which predicate(text) is True."""
    removed = 0
    body = doc.element.body
    for p in list(doc.paragraphs):
        text = p.text.strip()
        if not text or not predicate(text):
            continue
        body.remove(p._p)
        removed += 1
    return removed


def has_figure_caption(doc: Document, fig_num: int, *, snippet: str | None = None) -> bool:
    prefix = f"Figure {fig_num}."
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t.startswith(prefix):
            continue
        if snippet is None or snippet in t:
            return True
    return False


def paragraph_has_image(paragraph: Paragraph) -> bool:
    xml = paragraph._element.xml
    return "pic:pic" in xml or "w:drawing" in xml


def remove_old_threshold_figure(doc: Document) -> None:
    """Remove legacy Figura 3 caption and its adjacent standalone image, if present."""
    body = doc.element.body
    for p in list(doc.paragraphs):
        t = p.text.strip()
        is_legacy_caption = t.startswith("Figura 3. Threshold optimization") or (
            t.startswith(f"Figure {FIG_THRESHOLD}.")
            and "Threshold strategy comparison" in t
        )
        if not is_legacy_caption:
            continue
        prev = p._p.getprevious()
        if prev is not None:
            from docx.text.paragraph import Paragraph as DocxParagraph

            prev_para = DocxParagraph(prev, p._parent)
            if paragraph_has_image(prev_para) and not prev_para.text.strip():
                body.remove(prev)
        body.remove(p._p)


def ensure_section_48_body(doc: Document) -> None:
    if any("Feature importance analysis across five models" in p.text for p in doc.paragraphs):
        return
    h48 = find_para_contains(doc, "4.8. Feature Importance")
    insert_paragraph_after(
        h48,
        "Feature importance analysis across five models shows injection-phase variables account for "
        "24.6% of mean aggregated importance—the largest single process phase—followed by "
        "configuration/maintenance (17.6%). Distance-from-ideal-range features were the most "
        "important engineering category (28.6%). ROC analysis yielded micro-average AUC ≥ 0.9929 "
        "for all models (neural network: 0.9955).",
        "New paragraph",
    )


def fix_figure_numbering(doc: Document) -> None:
    """
    Renumber figures sequentially 1–11 and fix in-text citations.

    Intended order:
      1 workflow | 2 feature recall | 3 threshold trade-off | 4 dashboard
      5 confusion | 6 learning curves | 7–8 importance | 9–10 ROC | 11 calibration
    """
    remove_paragraphs_matching(
        doc,
        lambda t: t.startswith("Figure 4. Threshold strategy comparison")
        or t.startswith("Figure 5. Threshold strategy comparison"),
    )

    fig_replacements = [
        ("Figura 3. Threshold optimization for gas porosity defect.", THRESHOLD_CAPTION),
        ("Figura 3. Threshold optimization", THRESHOLD_CAPTION),
        (
            f"Figure {FIG_THRESHOLD}. Threshold strategy comparison (recall-first vs. F-beta β=2) "
            "for the neural network. for gas porosity defect.",
            THRESHOLD_CAPTION,
        ),
        (
            "Figure 5 illustrates the threshold analysis",
            f"Figure {FIG_THRESHOLD} compares four threshold strategies",
        ),
        ("Figure 5 compares four threshold strategies", f"Figure {FIG_THRESHOLD} compares four threshold strategies"),
        ("Figure 5. Threshold strategy comparison", f"Figure {FIG_THRESHOLD}. Threshold strategy comparison"),
        ("Figure 6 illustrates the dashboard", f"Figure {FIG_DASHBOARD} illustrates the dashboard"),
        ("Figure 6. Prototype Streamlit dashboard", f"Figure {FIG_DASHBOARD}. Prototype Streamlit dashboard"),
        ("Figure 6. Normalized confusion matrix", f"Figure {FIG_CONFUSION}. Normalized confusion matrix"),
        ("Figure 5. Normalized confusion matrix", f"Figure {FIG_CONFUSION}. Normalized confusion matrix"),
        (
            "Normalized confusion matrices (Figure 6)",
            f"Normalized confusion matrices (Figure {FIG_CONFUSION})",
        ),
        (
            "Calibration curves (Figure 12)",
            f"Calibration curves (Figure {FIG_CALIBRATION})",
        ),
        (
            "Calibration curves (Figure 10)",
            f"Calibration curves (Figure {FIG_CALIBRATION})",
        ),
    ]
    replace_in_all_paragraphs(doc, fig_replacements)

    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith(f"Figure {FIG_THRESHOLD}.") and "Threshold strategy comparison" in t:
            set_paragraph_text(p, THRESHOLD_CAPTION)


def insert_figure_after(
    paragraph: Paragraph,
    image_path: Path,
    caption: str,
    width_inches: float = 5.5,
) -> Paragraph:
    if not image_path.exists():
        print(f"  [AVISO] Figura não encontrada: {image_path}")
        return paragraph
    pic_para = insert_paragraph_after(paragraph)
    pic_para.add_run().add_picture(str(image_path), width=Inches(width_inches))
    cap = insert_paragraph_after(pic_para, caption, "Caption")
    cap.style = "Caption"
    return cap


def insert_figures_if_missing(doc: Document) -> None:
    if has_figure_caption(doc, FIG_CALIBRATION, snippet="Calibration curves"):
        return

    remove_old_threshold_figure(doc)

    if not has_figure_caption(doc, FIG_THRESHOLD, snippet="Threshold strategy"):
        thresh_anchor = None
        for snippet in (
            "Figure 3 compares four threshold strategies",
            "The probability distributions underlying these thresholds reveal",
        ):
            try:
                thresh_anchor = find_para_contains(doc, snippet)
                break
            except ValueError:
                continue
        if thresh_anchor is not None:
            insert_figure_after(
                thresh_anchor,
                FIGURES / "phase3_threshold_tradeoff.png",
                THRESHOLD_CAPTION,
            )

    if not has_figure_caption(doc, FIG_CONFUSION, snippet="Normalized confusion"):
        try:
            conf_anchor = find_para_contains(doc, "Figure 4. Prototype Streamlit dashboard")
        except ValueError:
            conf_anchor = find_para_contains(
                doc, "The prototype demonstrates that the trained model can be effectively integrated"
            )
        insert_figure_after(
            conf_anchor,
            FIGURES / "phase3_confusion_normalized_nn.png",
            f"Figure {FIG_CONFUSION}. Normalized confusion matrix (neural network, F-beta β=2 thresholds).",
        )

    if not has_figure_caption(doc, FIG_LEARNING, snippet="Learning curves"):
        body47 = find_para_contains(doc, "Formal overfitting analysis compared training-fold")
        insert_figure_after(
            body47,
            FIGURES / "phase1_learning_curves.png",
            f"Figure {FIG_LEARNING}. Learning curves — train vs. validation F1 (neural network, unified CV).",
        )

    ensure_section_48_body(doc)
    body48 = find_para_contains(doc, "Feature importance analysis across five models")

    fig_imp = body48
    if not has_figure_caption(doc, FIG_IMPORTANCE, snippet="Top-25"):
        fig_imp = insert_figure_after(
            body48,
            FIGURES / "phase3_feature_importance_top25.png",
            f"Figure {FIG_IMPORTANCE}. Top-25 feature importance (mean across five models).",
        )

    fig_phase = fig_imp
    if not has_figure_caption(doc, FIG_IMPORTANCE_PHASE, snippet="casting process phase"):
        fig_phase = insert_figure_after(
            fig_imp,
            FIGURES / "phase3_importance_by_phase.png",
            f"Figure {FIG_IMPORTANCE_PHASE}. Aggregated feature importance by casting process phase.",
        )

    fig_roc = fig_phase
    if not has_figure_caption(doc, FIG_ROC_TOP, snippet="ROC curves for the five"):
        fig_roc = insert_figure_after(
            fig_phase,
            FIGURES / "phase3_roc_curves_top_defects.png",
            f"Figure {FIG_ROC_TOP}. ROC curves for the five most frequent defect types (neural network).",
        )

    fig_roc_micro = fig_roc
    if not has_figure_caption(doc, FIG_ROC_MICRO, snippet="Micro-average ROC"):
        fig_roc_micro = insert_figure_after(
            fig_roc,
            FIGURES / "phase3_roc_micro_average.png",
            f"Figure {FIG_ROC_MICRO}. Micro-average ROC curves across all five model architectures.",
        )

    if not has_figure_caption(doc, FIG_CALIBRATION, snippet="Calibration curves"):
        insert_figure_after(
            fig_roc_micro,
            FIGURES / "phase3_calibration_by_model.png",
            f"Figure {FIG_CALIBRATION}. Calibration curves by model architecture (micro-average).",
        )


def ensure_figure_citations(doc: Document) -> None:
    """Add in-text references for figures 5–10 where the body discusses them."""
    replace_if_contains(
        doc,
        "The confusion matrix analysis for the most frequent defects demonstrates",
        f"The confusion matrix analysis for the most frequent defects (Figure {FIG_CONFUSION}) demonstrates",
    )
    replace_if_contains(
        doc,
        "Formal overfitting analysis compared training-fold, validation-fold, and test-set metrics",
        f"Formal overfitting analysis compared training-fold, validation-fold, and test-set metrics "
        f"(Figure {FIG_LEARNING})",
    )
    replace_if_contains(
        doc,
        "Feature importance analysis across five models shows injection-phase variables",
        f"Feature importance analysis across five models (Figures {FIG_IMPORTANCE}–{FIG_IMPORTANCE_PHASE}) "
        "shows injection-phase variables",
    )
    replace_if_contains(
        doc,
        "ROC analysis yielded micro-average AUC ≥ 0.9929 for all models (neural network: 0.9955), confirming",
        f"ROC analysis (Figures {FIG_ROC_TOP}–{FIG_ROC_MICRO}) yielded micro-average AUC ≥ 0.9929 for all models "
        "(neural network: 0.9955), confirming",
    )


def update_main_article(doc: Document) -> None:
    # --- Abstract & Introduction ---
    set_paragraph_text(
        find_para(doc, "Abstract:"),
        "Abstract: This study applies a comprehensive machine learning (ML) workflow for defect prediction "
        "in high-pressure die casting of aluminum alloys using a synthetic dataset grounded in expert knowledge. "
        "While traditional quality control methods tend to be reactive or limited in capturing complex "
        "interdependencies among process variables, ML offers a proactive, data-driven approach. This study "
        "implements an end-to-end workflow comprising seven integrated stages from problem framing to "
        "implementation and monitoring. The synthetic dataset encompasses 25,000 production records with "
        "15 process variables and 28 binary defect types. Class imbalance (94.40% defect-free) is addressed "
        "through cost-sensitive learning; an ablation study demonstrated that SMOTE oversampling did not "
        "improve performance when combined with cost-sensitive learning and was therefore excluded from the "
        "final pipeline. Feature engineering expanded the input space from 15 to 115 domain-informed features. "
        "Five model architectures were compared under a unified stratified five-fold cross-validation protocol "
        "(neural network, XGBoost, Random Forest, and two regularized logistic regression baselines). The "
        "neural network was selected for deployment based on the fastest inference (0.30 ms per 100 samples) "
        "and competitive cross-model performance (Table 8). The official operating policy applies F-beta (β=2) "
        f"per-defect thresholds, achieving recall {OFFICIAL['recall_micro']}, precision {OFFICIAL['precision_micro']}, "
        f"and F1-micro {OFFICIAL['f1_micro']} on the held-out test set. "
        "Feature importance analysis indicates that injection-phase variables account for 24.6% of aggregated "
        "importance across models, complemented by distance-from-ideal-range features. ROC analysis yielded "
        "micro-average AUC above 0.99 for all models. The study demonstrates that integrating the complete ML "
        "pipeline with rigorous cross-validation and statistical baselines enables effective multi-label defect "
        "prediction with emphasis on minimizing false negatives in manufacturing quality control.",
    )

    set_paragraph_text(
        find_para(doc, "The study achieves several key results."),
        "The study achieves several key results. Feature engineering expanded the feature space from 15 "
        "original process variables to 115 domain-informed features. Five architectures were evaluated under "
        "identical cross-validation and preprocessing protocols: PyTorch neural network, XGBoost, Random Forest, "
        "and regularized logistic regression (L2/Ridge and L1/Lasso). The neural network was selected for "
        "deployment based on the fastest inference (0.30 ms per 100 samples) and competitive F1-micro under "
        f"unified validation (Table 8: {TABLE8_NN['f1_micro']}). At the official F-beta (β=2) operating point, "
        f"the deployed model achieves recall {OFFICIAL['recall_micro']} and F1-micro {OFFICIAL['f1_micro']}. "
        "Logistic regression baselines achieved comparable F1-micro (~0.663) "
        "with substantially faster training, confirming that linear statistical models remain competitive "
        "with nonlinear approaches in this feature space.",
    )

    set_paragraph_text(
        find_para(doc, "This work makes two primary contributions."),
        "This work makes three primary contributions. First, it demonstrates that comprehensive implementation "
        "of the complete ML workflow enables effective, deployment-ready predictive systems for multi-label "
        "defect prediction. Second, it provides empirical evidence—through controlled ablation—that cost-sensitive "
        "learning alone is sufficient for class imbalance in this setting, and that SMOTE synthetic oversampling "
        "does not improve recall or F1 when cost-sensitive learning is active; on the contrary, SMOTE reduced "
        "recall by up to 19.7 percentage points (Random Forest) while increasing training time. Third, it extends "
        "the model comparison with interpretable statistical baselines (regularized logistic regression) and "
        "adds formal validation analyses: unified cross-validation, overfitting diagnostics, feature importance "
        "by process phase, ROC/AUC curves, and threshold strategy comparison.",
    )

    # --- Methodology 3.5 (preserve equation paragraphs; add revision notes) ---
    models_intro = find_para(doc, "Three model architectures were evaluated:")
    insert_note_after(
        models_intro,
        "Five model architectures were evaluated in the revised study: PyTorch neural network, XGBoost, "
        "Random Forest, and two regularized logistic regression baselines (L2/Ridge and L1/Lasso), all "
        "under an identical protocol (80/20 stratified split, unified five-fold CV, cost-sensitive learning, "
        "per-defect threshold optimization; SMOTE evaluated in ablation but excluded from the final pipeline).",
    )

    set_paragraph_text(
        find_para(doc, "The training pipeline applied to all three models"),
        "The training pipeline applied to all five models consists of four stages: feature scaling, "
        "cost-sensitive class weighting, stratified five-fold cross-validation on the development set, "
        "and per-defect threshold optimization on the held-out test set.",
    )

    set_paragraph_text(
        find_para(doc, "These parameters were then used to transform both training and test data."),
        "These parameters were then used to transform both training and test data. Stratified five-fold "
        "cross-validation was applied uniformly to all five model architectures on the development set "
        "(80% of data), stratifying by the binary indicator of any defect occurrence. Within each fold, "
        "the scaler was fit exclusively on the training partition; the validation partition was never "
        "used for fitting or oversampling. Cross-validation metrics (recall, precision, F1-micro, "
        "F1-macro) are reported as mean ± standard deviation across folds. The final model for test "
        "evaluation was trained on the complete development set.",
    )

    smote_heading = find_para(doc, "Class Balancing via SMOTE.")
    insert_note_after(
        smote_heading,
        "Implementation note (final pipeline): SMOTE was evaluated within cross-validation training folds "
        "only. An ablation comparing with/without SMOTE while keeping cost-sensitive learning active showed "
        "no F1 benefit and reduced recall; the deployed pipeline therefore uses cost-sensitive learning "
        "exclusively (Section 4.2; Supplementary Table S6). Equations (1)–(2) describe the SMOTE procedure "
        "as implemented in the ablation study.",
    )

    smote_detail = find_para_contains(doc, "In this multi-label problem, SMOTE is applied")
    insert_note_after(
        smote_detail,
        "Ablation summary: with cost-sensitive learning active, SMOTE reduced cross-validation recall "
        "(neural network −5.8 pp; XGBoost −12.0 pp; Random Forest −19.7 pp) without meaningful F1-micro "
        "gain and increased training time; SMOTE was excluded from the final configuration.",
    )

    thresh_heading = find_para(doc, "Adaptive Threshold Optimization.")
    insert_note_after(
        thresh_heading,
        "Official deployment policy: four strategies were compared on identical test-set probabilities "
        "(recall-first, F1-max, F-beta β=2, F-beta β=0.5). F-beta (β=2) was selected for the deployed "
        "neural network (recall 0.9828, precision 0.5068, F1-micro 0.6687; Section 4.4).",
    )

    set_paragraph_text(
        find_para(doc, "The threshold achieving highest recall"),
        "This threshold comparison was applied identically to all five model architectures using the "
        "same predicted probabilities from the final models, ensuring fair comparison of operating "
        "point strategies without retraining.",
    )

    # 3.5.5 Logistic Regression
    note_para = find_para(doc, "Note on Hyperparameter Selection.")
    if not any(p.text.startswith("3.5.5.") for p in doc.paragraphs):
        h = insert_paragraph_before(note_para, "3.5.5. Logistic Regression Baselines", "Heading 1")
        insert_paragraph_before(
            note_para,
            "Two regularized logistic regression models were included as interpretable statistical baselines, "
            "implemented via one-vs-rest classifiers with class_weight='balanced'. L2 (Ridge) regularization "
            "uses the L-BFGS solver (C=1.0, max_iter=2000); L1 (Lasso) uses the liblinear solver for sparse "
            "coefficient selection. Standardized features ensure coefficient magnitudes are directly "
            "comparable as importance measures.",
            "New paragraph",
        )
        # move new paragraph to right after heading
        # (insert_paragraph_before note puts both before note; order is: h, body, note - OK)

    # --- Section 4.2 ---
    set_paragraph_text(find_para(doc, "4.2. Data Balancing Results"), "4.2. Class Imbalance Handling Results")

    set_paragraph_text(
        find_para(doc, "The class balancing strategy addressed the severe imbalance"),
        "The class balancing strategy addressed the severe imbalance characteristic of the dataset, "
        f"where only 5.60% of samples (1,401 of 25,000) exhibited any defect. Table {TABLE_SMOTE} presents the "
        "original frequency distribution for the most common defects in the training set. Individual "
        "defect frequencies ranged from 3.22% (gas porosity) to 0.00% (ejector pin marks). This "
        "extreme imbalance poses a significant challenge, as standard optimization would favor "
        "predicting the majority class almost exclusively.",
    )

    set_paragraph_text(
        find_para_contains(doc, "Table 9 illustrates the dramatic effect of SMOTE"),
        "An ablation study compared two conditions under identical five-fold cross-validation: "
        "(a) cost-sensitive learning only, and (b) SMOTE oversampling within training folds combined "
        f"with cost-sensitive learning. Table {TABLE_SMOTE} retains the original SMOTE frequency projections for "
        "reference; however, the ablation results (Table S6, Supplementary Materials) demonstrate "
        "that SMOTE did not improve performance when cost-sensitive learning was active.",
    )

    set_paragraph_text(
        find_para(doc, "SMOTE application to the training data increased"),
        "Without SMOTE, cross-validation recall was 0.9511 ± 0.0098 (neural network), 0.8576 ± 0.0123 "
        "(XGBoost), and 0.8204 ± 0.0114 (Random Forest). With SMOTE, recall dropped to 0.8928, 0.7377, "
        "and 0.6239 respectively. F1-micro changes were negligible or negative (average ΔF1 = −0.015). "
        "SMOTE also inflated the train-validation performance gap, suggesting memorization of synthetic "
        "samples in the 115-dimensional feature space. Based on these results, the final pipeline uses "
        "cost-sensitive learning exclusively.",
    )

    set_paragraph_text(
        find_para(doc, "The class weights incorporated into the loss function"),
        "Cost-sensitive mechanisms were: class-weighted loss for the neural network (weights 1.0–10.0), "
        "class_weight='balanced' for Random Forest and logistic regression, and scale_pos_weight per "
        "defect for XGBoost. At the official F-beta (β=2) operating point, the deployed neural network "
        f"achieves test-set recall of {OFFICIAL['recall_micro']} without synthetic oversampling.",
    )

    # --- Section 4.3 ---
    set_paragraph_text(
        find_para(doc, "Three model architectures were evaluated using the same data pipeline"),
        "Five model architectures were evaluated under an identical protocol: 80/20 stratified split, "
        "Z-score normalization, cost-sensitive learning (no SMOTE), unified stratified five-fold "
        "cross-validation on the development set, and per-defect threshold optimization. Results on "
        "the held-out test set (n = 5,000) are summarized in Table 8.",
    )

    set_paragraph_text(
        find_para(doc, "XGBoost training completed in 53.67 seconds"),
        "XGBoost training completed in 23.4 seconds with scale_pos_weight per defect. Cross-validation "
        "recall was 0.8576 ± 0.0123 and F1-micro 0.6692 ± 0.0100. A substantial train-validation gap "
        "in F1 (0.23) indicates overfitting tendency, although test-set performance remained stable.",
    )

    set_paragraph_text(
        find_para(doc, "Random Forest training required 210.26 seconds"),
        "Random Forest training required 55.8 seconds. It achieved the highest test recall (0.9968) "
        "but the lowest F1-micro (0.5847) due to very low precision (0.4137), and inference latency "
        "of 1,269 ms per 100 samples—impractical for real-time deployment.",
    )

    set_paragraph_text(
        find_para(doc, "Stratified five-fold cross-validation was performed only for the neural network"),
        "Stratified five-fold cross-validation was applied uniformly to all five architectures. "
        "Neural network cross-validation recall was 0.9511 ± 0.0098 (F1-micro 0.6657 ± 0.0071). "
        "Logistic regression L1 achieved the highest cross-validation F1-micro (0.6730 ± 0.0084), "
        "while L2/Ridge offered the fastest training (11.3 s) with competitive test metrics.",
    )

    update_table_18(doc.tables[18])
    update_table_17_smote_reference(doc.tables[17])
    update_table_19_thresholds(doc.tables[19])
    update_table_20_confusion(doc.tables[20])

    # Corrigir numeração duplicada da tabela de thresholds
    thresh_cap = find_para_contains(doc, "Optimized Thresholds and Performance for Selected Defects")
    set_paragraph_text(
        thresh_cap,
        f"Table {TABLE_THRESHOLDS}. Optimized Thresholds and Performance for Selected Defects "
        "(F-beta β=2 strategy, neural network, test set n = 5,000).",
    )
    smote_cap = find_para_contains(doc, "Defect Frequencies and SMOTE Balancing Effect")
    set_paragraph_text(
        smote_cap,
        f"Table {TABLE_SMOTE}. Defect Frequencies and SMOTE Balancing Effect (Training Set, n = 20,000). "
        "Reference distribution from ablation study; SMOTE was not used in the final pipeline.",
    )

    set_paragraph_text(
        find_para(doc, "As shown in Table 8, XGBoost achieved the highest F1-micro"),
        "As shown in Table 8, all five models achieved comparable F1-micro (0.5847–0.6639). Logistic "
        "regression L1 achieved the highest cross-validation F1-micro (0.6730), while the neural network "
        "achieved the highest test recall (0.9930). The narrow F1 range across architectures suggests "
        "that domain-informed feature engineering contributes more to performance than algorithm choice alone.",
    )

    set_paragraph_text(
        find_para(doc, "Random Forest achieved the highest recall (0.9838)"),
        "Random Forest achieved the highest recall (0.9968), followed by the neural network (0.9930) "
        "and logistic regression models (~0.9844). XGBoost recall was 0.9811.",
    )

    set_paragraph_text(
        find_para(doc, "The inference time differences are equally significant"),
        "Inference time differences are decisive for real-time deployment (Table 8). The neural network "
        "processes 100 samples in 0.30 ms; logistic regression requires 2.77–2.85 ms; XGBoost 16.96 ms; "
        "and Random Forest 1,269 ms. The neural network's inference is 57× faster than XGBoost.",
    )

    set_paragraph_text(
        find_para_contains(doc, "Table 8. Model Comparison"),
        "Table 8. Model Comparison on Test Set (n = 5,000). Metrics obtained with per-defect recall-first "
        "thresholds for fair cross-model comparison; the deployed neural network uses F-beta (β=2) thresholds "
        "(Section 4.4).",
    )

    set_paragraph_text(
        find_para(doc, "The PyTorch neural network was selected for production deployment"),
        "The PyTorch neural network was selected for production deployment based on the fastest inference "
        f"({TABLE8_NN['inference_ms']} ms per 100 samples), competitive F1-micro ({TABLE8_NN['f1_micro']}), "
        f"and suitability for real-time CPPS integration. Although Random Forest achieved marginally higher "
        "recall under recall-first thresholds (0.9968), its inference latency (1,269 ms per 100 samples) exceeds "
        "real-time requirements by three orders of magnitude. The official deployed operating point uses "
        f"F-beta (β=2) thresholds (recall {OFFICIAL['recall_micro']}, precision {OFFICIAL['precision_micro']}).",
    )

    if not any("Four threshold strategies were compared" in p.text for p in doc.paragraphs):
        thresh_para = find_para(doc, "Per-defect threshold optimization produced thresholds")
        insert_paragraph_before(
            thresh_para,
            "Four threshold strategies were compared on identical test-set probabilities (Section 3.5.1). "
            "Recall-only optimization achieved recall 0.9925 but precision 0.4876 for the neural network. "
            "The selected F-beta (β=2) strategy achieved recall 0.9828, precision 0.5068, and F1-micro "
            "0.6687—maintaining near-maximum defect capture while modestly improving precision. "
            "Full comparison in Supplementary Materials (Section 4.4).",
            "New paragraph",
        )

    set_paragraph_text(
        find_para(doc, "Per-defect threshold optimization produced thresholds"),
        "Per-defect threshold optimization under the F-beta (β=2) strategy produced thresholds "
        f"ranging from 0.05 to 0.56 across defect types (Lipton et al., 2014). Table {TABLE_THRESHOLDS} presents "
        "the optimized thresholds for the most frequent defects. For the most frequent defects, "
        "optimal thresholds were: gas porosity (0.05), density deviation (0.10), and cold shut "
        "(0.21). Lower thresholds for frequent defects reflect high predicted probabilities for "
        "true positives; higher thresholds for rare defects (e.g., shrinkage porosity at 0.54) "
        "balance recall against false alarms.",
    )

    set_paragraph_text(
        find_para(doc, "For less frequent defects, optimal thresholds were often lower still"),
        "For less frequent defects, optimal thresholds varied with predicted probability "
        "distributions. Low tensile strength achieved an optimal threshold of 0.56, warpage 0.53, "
        "and shrinkage porosity 0.54. These values reflect the F-beta (β=2) compromise between "
        "recall prioritization and precision improvement relative to recall-only optimization.",
    )

    set_paragraph_text(
        find_para(doc, "The confusion matrix analysis for the most frequent defects"),
        "The confusion matrix analysis for the most frequent defects demonstrates the trade-off "
        f"between recall and precision that characterizes the optimized model. Table {TABLE_CONFUSION} presents "
        "confusion matrix metrics for the eight most frequent defect types on the test set. "
        "For gas porosity (141 positive instances), the model detected all 141 defects "
        "(recall 1.000) with 120 false positives (precision 0.540).",
    )

    set_paragraph_text(
        find_para(doc, "For cold shut (138 occurrences in test set)"),
        "For cold shut (138 occurrences in test set), the model achieved perfect recall (1.000), "
        "identifying all 138 defects with zero false negatives. Among 234 samples predicted as "
        "cold shut defective, 138 were correct and 96 were false positives (precision 0.590).",
    )

    set_paragraph_text(
        find_para(doc, "For incomplete fill (105 occurrences in test set)"),
        "For incomplete fill (105 occurrences in test set), the model detected 103 defects "
        "(recall 0.981), missing two. Of 198 samples classified as defective, 103 were correct "
        "and 95 were false positives (precision 0.520).",
    )

    set_paragraph_text(
        find_para(doc, "For less frequent defects, the pattern"),
        "For less frequent defects, the pattern of high recall with moderate precision persists. "
        "Gas bubbles (129 occurrences in the test set) achieved recall 0.977 with precision 0.624. "
        "The final model achieves micro-averaged recall of 0.983 and precision of 0.507 across "
        "all 28 defect types under the F-beta (β=2) operating policy.",
    )

    # --- New sections 4.7 and 4.8 BEFORE Discussion (insert in reverse order) ---
    discussion = find_para(doc, "5. Discussion")
    if not any(p.text.startswith("4.7.") for p in doc.paragraphs):
        h47 = insert_paragraph_before(
            discussion, "4.7. Cross-Validation, Overfitting, and SMOTE Ablation", "Heading 1"
        )
        insert_paragraph_after(
            h47,
            "Formal overfitting analysis compared training-fold, validation-fold, and test-set metrics "
            "(threshold 0.5) under the unified CV protocol. Validation-test gaps were negligible "
            "(|Δ| < 0.03). Train-validation gaps were larger for tree-based models (Random Forest "
            "F1 gap 0.30; XGBoost 0.23; neural network 0.05). SMOTE inflated train-test gaps, "
            "supporting its exclusion (Table S6–S7, Supplementary Materials).",
            "New paragraph",
        )
        h48 = insert_paragraph_before(discussion, "4.8. Feature Importance and ROC Analysis", "Heading 1")
        insert_paragraph_after(
            h48,
            "Feature importance analysis across five models shows injection-phase variables account for "
            "24.6% of mean aggregated importance—the largest single process phase—followed by "
            "configuration/maintenance (17.6%). Distance-from-ideal-range features were the most "
            "important engineering category (28.6%). ROC analysis yielded micro-average AUC ≥ 0.9929 "
            "for all models (neural network: 0.9955).",
            "New paragraph",
        )

    fix_remaining_contradictions(doc)

    fix_table_numbering(doc)
    insert_figures_if_missing(doc)
    ensure_figure_citations(doc)
    fix_figure_numbering(doc)


def fix_remaining_contradictions(doc: Document) -> None:
    """Substitui trechos legados ainda presentes no .docx original copiado."""

    replace_if_contains(
        doc,
        "Note on Hyperparameter Selection.",
        "Note on Hyperparameter Selection. For all five model architectures, hyperparameters were set "
        "based on established best practices and preliminary experimentation rather than systematic "
        "optimization (e.g., grid search or Bayesian optimization). Tables 4–6 present the fixed values "
        "used throughout the experiments. While systematic hyperparameter tuning could potentially improve "
        "model performance, the focus of this study was on demonstrating the complete workflow and "
        "responding to banca feedback through unified validation rather than achieving optimal accuracy "
        "for any single architecture.",
    )

    replace_if_contains(
        doc,
        "The trained model is packaged into a comprehensive artifact",
        "The trained model is packaged into a comprehensive artifact including the neural network, feature "
        "scaler, F-beta (β=2) optimized thresholds, and all necessary components for inference. This artifact "
        "can be integrated into cyber-physical production systems (CPPS), Manufacturing Execution Systems "
        "(MES), or real-time control dashboards. A key requirement is real-time detection capability; the "
        "model must generate inferences within milliseconds to support immediate corrective action. The "
        f"neural network's inference time ({TABLE8_NN['inference_ms']} ms per 100 samples) is well within "
        "the requirements for typical die casting cycle times of 30 seconds to 2 minutes.",
    )

    replace_if_contains(
        doc,
        "The probability distributions underlying these thresholds reveal",
        "Figure 3 compares four threshold strategies on identical test-set probabilities. Under the selected "
        "F-beta (β=2) policy, gas porosity uses a threshold of 0.05 (recall 1.000, precision 0.540), "
        "reflecting the model's ability to assign high probabilities to true positive cases while moderating "
        "false alarms relative to recall-only optimization (recall-only: recall 0.9925, precision 0.4876). "
        "For well-represented defects, non-defective "
        "samples concentrate at low probabilities while defective samples shift toward higher values, enabling "
        "per-defect threshold placement on the ROC-informed probability scale.",
    )

    replace_if_contains(
        doc,
        "Figura 3. Threshold optimization",
        THRESHOLD_CAPTION,
    )

    replace_if_contains(
        doc,
        "For rare defects, the distributions overlap more substantially",
        "For rare defects, predicted probability distributions overlap more substantially, requiring "
        "higher per-defect thresholds under F-beta (β=2) (e.g., shrinkage porosity at 0.54, warpage at 0.53) "
        "to balance recall against false positives. This differs from a fixed 0.5 cut-off, which would "
        "substantially reduce recall for rare defects. Per-defect optimization ensures each defect type "
        "receives an operating point aligned with its frequency and the model's discriminative capability, "
        "as validated by ROC/AUC analysis (Section 4.8) and predicted-versus-actual calibration (Section 4.9).",
    )

    replace_if_contains(
        doc,
        f"The overall row in Table {TABLE_CONFUSION} presents micro-averaged metrics",
        f"The overall row in Table {TABLE_CONFUSION} presents micro-averaged metrics aggregated across all 28 defect types "
        f"under F-beta (β=2) thresholds: recall {OFFICIAL['recall_micro']}, precision {OFFICIAL['precision_micro']}, "
        f"F1-micro {OFFICIAL['f1_micro']}, and macro-averaged F1 {OFFICIAL['f1_macro']}. Micro-averaging treats "
        "each prediction as an independent binary classification, summing true positives, false positives, and "
        "false negatives across all defects before computing precision and recall.",
    )

    replace_if_contains(
        doc,
        "The pattern is consistent",
        "The pattern is consistent across defect types: the model operates under an F-beta (β=2) policy that "
        "prioritizes recall while improving precision relative to recall-only optimization. Precision values "
        "reflect class imbalance and the operational priority of detecting defects over avoiding false alarms. "
        f"The final deployed model achieves micro-averaged recall of {OFFICIAL['recall_micro']}, micro-averaged "
        f"precision of {OFFICIAL['precision_micro']}, micro-averaged F1 of {OFFICIAL['f1_micro']}, and "
        f"macro-averaged F1 of {OFFICIAL['f1_macro']}. The difference between micro and macro F1 (0.054) "
        "indicates that performance is somewhat better for frequent defects than for rare ones.",
    )

    replace_if_contains(
        doc,
        "The third advance emerges from the comparative evaluation of neural networks, XGBoost, and Random Forest",
        "The third advance emerges from the comparative evaluation of five architectures—neural network, "
        "XGBoost, Random Forest, and two regularized logistic regression baselines—under an identical "
        "protocol (stratified 80/20 split, unified five-fold cross-validation, cost-sensitive learning, "
        "no SMOTE). This provides evidence on trade-offs between recall, precision, inference latency, and "
        "training cost in multi-label defect prediction. The neural network was selected for deployment based "
        f"on sub-millisecond inference ({TABLE8_NN['inference_ms']} ms per 100 samples) and operational "
        "suitability, while the official F-beta (β=2) operating point balances recall and precision on the "
        "held-out test set.",
    )

    replace_if_contains(
        doc,
        "The model's recall of 97.63% ensures",
        "The model's recall of 98.28% (F-beta β=2 operating point) ensures",
    )

    replace_if_contains(
        doc,
        "achieving 97.63% recall, inspection effort",
        f"achieving {OFFICIAL['recall_micro']} recall (F-beta β=2), inspection effort",
    )

    replace_if_contains(
        doc,
        "The 2.37% of defects missed by the model",
        "The 1.72% of defects missed by the model",
    )

    replace_if_contains(
        doc,
        "The threshold optimization strategy, which maximizes recall per defect type",
        "The threshold optimization strategy, based on F-beta (β=2) per defect type",
    )

    replace_if_contains(
        doc,
        "Organizations with higher costs of field failure relative to inspection costs would select lower thresholds to maximize recall",
        "Organizations with higher costs of field failure relative to inspection costs may select lower thresholds "
        "(closer to recall-first optimization), while those prioritizing inspection capacity may shift toward "
        "F-beta (β=0.5) or F1-maximization strategies documented in the supplementary materials",
    )

    replace_if_contains(
        doc,
        "The model's inference time of 0.33 ms per 100 samples",
        f"The model's inference time of {TABLE8_NN['inference_ms']} ms per 100 samples",
    )

    replace_if_contains(
        doc,
        "Fourth, the precision values around 0.50-0.54 indicate",
        f"Fourth, the precision values around {OFFICIAL['precision_micro']} (micro-average under F-beta β=2) indicate",
    )

    replace_if_contains(
        doc,
        "A multi-layer perceptron with cost-sensitive loss achieved the highest recall (99.30%)",
        f"A multi-layer perceptron with cost-sensitive loss was selected for deployment (inference "
        f"{TABLE8_NN['inference_ms']} ms per 100 samples; official operating recall {OFFICIAL['recall_micro']} "
        f"under F-beta β=2). Regularized logistic regression achieved comparable F1-micro (~0.663). "
        f"SMOTE was evaluated but excluded from the final pipeline based on ablation evidence.",
    )

    # Expandir §4.8 (ROC + ligação com threshold) se já existir
    replace_if_contains(
        doc,
        "ROC analysis yielded micro-average AUC ≥ 0.9929",
        "ROC analysis yielded micro-average AUC ≥ 0.9929 for all models (neural network: 0.9955), confirming "
        "strong discriminative ability independent of the chosen operating threshold. High AUC supports "
        "per-defect threshold tuning on the same probability scale used in Section 4.4, rather than indicating "
        "that a fixed 0.5 cut-off would be appropriate for imbalanced multi-label defect detection",
    )

    # §4.9 Previsto × realizado
    if not any(p.text.startswith("4.9.") for p in doc.paragraphs):
        discussion = find_para(doc, "5. Discussion")
        h49 = insert_paragraph_before(discussion, "4.9. Predicted versus Actual Validation", "Heading 1")
        insert_paragraph_after(
            h49,
            "Predicted-versus-actual validation complements threshold and ROC analysis by comparing predicted "
            "probabilities with observed defect frequencies on the held-out test set. Calibration curves "
            "(Figure 11) plot mean predicted probability against the fraction of actual positives in each "
            "probability bin, aggregated across all 28 defect types; points near the diagonal indicate "
            f"well-calibrated probabilities. Normalized confusion matrices (Figure {FIG_CONFUSION}) provide a per-defect "
            "view of predicted versus actual class assignments under F-beta (β=2) thresholds. Together, "
            "these diagnostics confirm that the model's probability outputs are usable for operational "
            "threshold selection and that classification errors concentrate in false positives rather than "
            "false negatives under the selected operating policy.",
            "New paragraph",
        )

    # --- Discussion updates ---
    set_paragraph_text(
        find_para(doc, "The class imbalance problem manifests primarily through false negatives, which occur"),
        "The class imbalance problem manifests primarily through false negatives. This study addresses "
        "imbalance through cost-sensitive learning, with ablation evidence that SMOTE does not improve—"
        "and often harms—recall when cost-sensitive mechanisms are active, revising the initial "
        "hypothesis that SMOTE and cost-sensitive learning together outperform either technique alone.",
    )

    set_paragraph_text(
        find_para(doc, "The second advance involves the integration of SMOTE oversampling"),
        "The second advance concerns empirical validation of class imbalance strategies through "
        "controlled ablation. The side-by-side comparison provides reproducible evidence that synthetic "
        "oversampling in high-dimensional multi-label settings can reduce recall and increase overfitting "
        "without F1 benefit when cost-sensitive learning is already active.",
    )

    set_paragraph_text(
        find_para(doc, "Second, methodological differences in the model comparison"),
        "Second, following banca feedback, the model comparison was revised to employ identical "
        "stratified five-fold cross-validation for all architectures, unified cost-sensitive learning "
        "(including scale_pos_weight for XGBoost), exclusion of SMOTE based on ablation evidence, "
        "and addition of two regularized logistic regression baselines.",
    )

    set_paragraph_text(
        find_para_contains(doc, "Domain-informed feature engineering expanded the representation from 15 raw variables"),
        "Domain-informed feature engineering expanded the representation from 15 raw variables to 115 "
        "features. Five architectures were compared under unified cross-validation; a multi-layer "
        f"perceptron with cost-sensitive loss was selected for deployment (inference {TABLE8_NN['inference_ms']} "
        f"ms per 100 samples; official operating recall {OFFICIAL['recall_micro']} under F-beta β=2). "
        "Regularized logistic regression achieved comparable F1-micro (~0.663). "
        "SMOTE was evaluated but excluded from the final pipeline based on ablation evidence.",
    )

    set_paragraph_text(
        find_para(doc, "The theoretical contributions include domain-informed feature engineering, combined class balancing"),
        "The theoretical contributions include domain-informed feature engineering, empirically validated "
        "cost-sensitive class imbalance handling (with evidence against default SMOTE), unified "
        "cross-validation, statistical baselines via regularized logistic regression, and F-beta "
        "threshold optimization. The practical contributions encompass sub-millisecond inference, "
        "feature importance by process phase, ROC/AUC validation, and inspection screening.",
    )


def update_supplementary(doc: Document) -> None:
    doc.add_page_break()
    doc.add_paragraph("4. Unified Cross-Validation and Validation Analyses", style="Heading 1")
    doc.add_paragraph(
        "Following banca feedback, all model comparisons were repeated under a unified experimental "
        "protocol: stratified 80/20 split (random_state=42), stratified five-fold cross-validation "
        "on the development set, cost-sensitive learning for all architectures, and SMOTE applied "
        "only within CV training folds when evaluating the oversampling ablation. The final pipeline "
        "excludes SMOTE based on ablation results (Table S6).",
        style="New paragraph",
    )

    doc.add_paragraph("4.1. Cross-Validation and SMOTE Ablation", style="Heading 1")
    doc.add_paragraph("Table S6. SMOTE Ablation — Cross-Validation Metrics (Mean ± Std, 5 Folds).", style="Table title")

    t6 = doc.add_table(rows=7, cols=7)
    t6.style = "Table Grid"
    smote_rows = [
        ["Model", "Condition", "Recall (CV)", "Precision (CV)", "F1-micro (CV)", "F1-macro (CV)", "CV Time (s)"],
        ["PyTorch NN", "No SMOTE", "0.9511 ± 0.0098", "0.5122 ± 0.0077", "0.6657 ± 0.0071", "0.6044 ± 0.0114", "417.4"],
        ["PyTorch NN", "With SMOTE", "0.8928 ± 0.0102", "0.5389 ± 0.0105", "0.6721 ± 0.0106", "0.6093 ± 0.0147", "366.4"],
        ["XGBoost", "No SMOTE", "0.8576 ± 0.0123", "0.5488 ± 0.0107", "0.6692 ± 0.0100", "0.6045 ± 0.0129", "120.2"],
        ["XGBoost", "With SMOTE", "0.7377 ± 0.0088", "0.5987 ± 0.0079", "0.6609 ± 0.0075", "0.5809 ± 0.0107", "329.9"],
        ["Random Forest", "No SMOTE", "0.8204 ± 0.0114", "0.5491 ± 0.0151", "0.6578 ± 0.0138", "0.5663 ± 0.0313", "204.5"],
        ["Random Forest", "With SMOTE", "0.6239 ± 0.0052", "0.6059 ± 0.0077", "0.6147 ± 0.0049", "0.5207 ± 0.0152", "949.0"],
    ]
    fill_table(t6, smote_rows)

    doc.add_paragraph("4.2. Overfitting Analysis", style="Heading 1")
    doc.add_paragraph("Table S7. Overfitting Diagnostics (threshold 0.5).", style="Table title")
    t7 = doc.add_table(rows=4, cols=6)
    t7.style = "Table Grid"
    fill_table(
        t7,
        [
            ["Model", "Metric", "Train (CV)", "Validation (CV)", "Test", "Gap Train-Val"],
            ["PyTorch NN", "F1", "0.7176 ± 0.0021", "0.6721 ± 0.0106", "0.6701", "0.0455"],
            ["XGBoost", "F1", "0.8908 ± 0.0018", "0.6609 ± 0.0075", "0.6629", "0.2299"],
            ["Random Forest", "F1", "0.9135 ± 0.0026", "0.6147 ± 0.0049", "0.6215", "0.2987"],
        ],
    )

    doc.add_paragraph("4.3. Extended Model Comparison (Five Architectures)", style="Heading 1")
    doc.add_paragraph("Table S8. Model Comparison on Test Set (No SMOTE, Cost-Sensitive).", style="Table title")
    t8 = doc.add_table(rows=6, cols=7)
    t8.style = "Table Grid"
    fill_table(
        t8,
        [
            ["Model", "F1-Micro", "F1-Macro", "Precision", "Recall", "Training (s)", "Inference (ms/100)"],
            ["PyTorch NN", "0.6639", "0.6088", "0.4986", "0.9930", "48.5", "0.30"],
            ["XGBoost", "0.6610", "0.6200", "0.4984", "0.9811", "23.4", "16.96"],
            ["Random Forest", "0.5847", "0.5600", "0.4137", "0.9968", "55.8", "1269.45"],
            ["Logistic Reg. (L2)", "0.6630", "0.6206", "0.4999", "0.9844", "11.3", "2.85"],
            ["Logistic Reg. (L1)", "0.6636", "0.6191", "0.5005", "0.9844", "510.2", "2.77"],
        ],
    )

    doc.add_paragraph("4.4. Threshold Strategy Comparison", style="Heading 1")
    doc.add_paragraph(
        "Four per-defect threshold strategies were evaluated. The F-beta (β=2) strategy was selected "
        "for the deployed neural network (recall 0.9828, precision 0.5068, F1-micro 0.6687).",
        style="New paragraph",
    )

    doc.add_paragraph("4.5. Feature Importance by Process Phase", style="Heading 1")
    doc.add_paragraph("Table S9. Aggregated Feature Importance by Casting Phase.", style="Table title")
    t9 = doc.add_table(rows=6, cols=3)
    t9.style = "Table Grid"
    fill_table(
        t9,
        [
            ["Process Phase", "Aggregated Importance", "N Features"],
            ["Global/Aggregation", "0.3012", "4"],
            ["Injection", "0.2459", "36"],
            ["Configuration/Maintenance", "0.1758", "35"],
            ["Intensification", "0.1422", "16"],
            ["Cooling", "0.1072", "14"],
        ],
    )

    doc.add_paragraph("4.6. ROC/AUC Summary", style="Heading 1")
    doc.add_paragraph("Table S10. Micro-Average AUC on Test Set.", style="Table title")
    t10 = doc.add_table(rows=6, cols=2)
    t10.style = "Table Grid"
    fill_table(
        t10,
        [
            ["Model", "AUC (micro-average)"],
            ["PyTorch NN", "0.9955"],
            ["XGBoost", "0.9952"],
            ["Random Forest", "0.9945"],
            ["Logistic Reg. (L2)", "0.9929"],
            ["Logistic Reg. (L1)", "0.9932"],
        ],
    )

    doc.add_paragraph("4.7. Supplementary Figures", style="Heading 1")
    fig_anchor = doc.add_paragraph(
        "The following figures support the validation analyses reported in the main manuscript.",
        style="New paragraph",
    )
    supp_figures = [
        (FIGURES / "phase1_learning_curves.png", "Figure S1. Learning curves (neural network, unified CV)."),
        (FIGURES / "phase3_threshold_tradeoff.png", "Figure S2. Threshold strategy comparison."),
        (FIGURES / "phase3_feature_importance_top25.png", "Figure S3. Top-25 feature importance."),
        (FIGURES / "phase3_importance_by_phase.png", "Figure S4. Feature importance by process phase."),
        (FIGURES / "phase3_roc_curves_top_defects.png", "Figure S5. ROC curves for top defects."),
        (FIGURES / "phase3_roc_micro_average.png", "Figure S6. Micro-average ROC curves."),
        (FIGURES / "phase3_calibration_by_model.png", "Figure S7. Calibration curves by model."),
        (FIGURES / "phase3_confusion_normalized_nn.png", "Figure S8. Normalized confusion matrix (neural network)."),
    ]
    anchor = fig_anchor
    for path, caption in supp_figures:
        anchor = insert_figure_after(anchor, path, caption, width_inches=5.0)


STALE_ARTICLE_VARIANTS = [
    BASE / "Artigo 2 - FINAL - BANCA REVISADO - TABELAS.docx",
    BASE / "Artigo 2 - FINAL - BANCA REVISADO - FIGURAS.docx",
    BASE / "_table_fix_test.docx",
]


def remove_stale_variants() -> None:
    for path in STALE_ARTICLE_VARIANTS:
        if path.exists():
            path.unlink()
            print(f"[LIMPO] Removido arquivo intermediário: {path.name}")


def main() -> None:
    shutil.copy2(MAIN_SRC, MAIN_OUT)
    shutil.copy2(SUPP_SRC, SUPP_OUT)

    print(f"[*] Gerando arquivo final: {MAIN_OUT.name}...")
    main_doc = Document(str(MAIN_OUT))
    update_main_article(main_doc)
    main_doc.save(str(MAIN_OUT))
    print(f"[OK] {MAIN_OUT}")

    print(f"[*] Atualizando suplementar: {SUPP_OUT.name}...")
    supp_doc = Document(str(SUPP_OUT))
    update_supplementary(supp_doc)
    supp_doc.save(str(SUPP_OUT))
    print(f"[OK] {SUPP_OUT}")

    remove_stale_variants()
    print(f"\nArquivo final do artigo: {MAIN_OUT}")
    print("Originais preservados. Abra o arquivo final no Word para revisão visual.")


if __name__ == "__main__":
    main()
